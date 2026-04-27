import os
from typing import List
from ..models import ExtractResult, MAX_FULL_TEXT, MAX_TRUNCATED_TEXT


class DocxExtractor:
    def supported_extensions(self) -> List[str]:
        return ["docx"]

    def extract(self, file_path: str) -> ExtractResult:
        try:
            from docx import Document
        except ImportError:
            return ExtractResult.failed("未安装 python-docx 库，无法提取 .docx 文件")

        try:
            doc = Document(file_path)
        except Exception as e:
            return ExtractResult.failed(f"无法打开 docx 文件: {e}")

        parts = []
        paragraph_count = 0
        table_count = 0

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
                paragraph_count += 1

        for table in doc.tables:
            table_parts = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_parts.append("| " + " | ".join(cells) + " |")
            if table_parts:
                parts.append("\n".join(table_parts))
                table_count += 1

        text = "\n\n".join(parts)
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

        if len(text) <= MAX_FULL_TEXT:
            return ExtractResult.success(text, method="full", metadata={
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "char_count": len(text),
                "size_bytes": file_size,
            })

        truncated = text[:MAX_TRUNCATED_TEXT]
        return ExtractResult.success(truncated, method="truncated", metadata={
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "char_count": len(text),
            "displayed_chars": MAX_TRUNCATED_TEXT,
            "size_bytes": file_size,
        })
